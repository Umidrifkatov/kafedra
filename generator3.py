import docx
import os
from heading import main_head
from students import students_list
from quest_list3 import quest
from thname3 import themas
  
doc = docx.Document()
  


theme =int(input('\n\nВыбор темы:\n\n  1 - Заболеваемость и смертность от рака молочной железы  \n  2 - Принципы системной терапии рака молочной железы. Химиотерапия. Гормонотерапия рака молочной железы  \n  3 - Доброкачественная патология молочной железы. Принципы диагностики и лечения. \n  4 - «Хирургическое лечение рака молочной железы. Послеоперационные осложнения. Виды реконструктивных операций при раке молочной железы. \n  5 - Стадирование рака молочной железы на основании клинических и инструментальных данных. Современная патоморфологическая  классификация рака молочной железы. Редкие формы рака молочной железы. Соотношение клинических признаков заболевания с международной классификацией  рака молочной железы.  \n  6 - Течение и прогноз рака молочной железы. Наследственный рак молочной железы. Основные генетические мутации при раке молочной железы. \n 7 - Верификация диагноза рак молочной железы. биопсия \n 8 - Лучевая терапия рака молочной железы \n 9 - Зачет \n 10 - Реабилитация \n\nВедите числовое значение от 1 до 6.\nОжидание выбора темы .........\n'))
per = 0
print(len(quest[theme]))
q = quest[theme]

for i in students_list:
    student_name = i[0]
    st_group = i[1]

    header = doc.add_paragraph(main_head)
    header.alignment = 1

    group = doc.add_paragraph('Группа - ')
    group.add_run(st_group).bold = True
    group.add_run('         ФИО студента - ')
    group.add_run(student_name.upper()).bold = True

    theme_r = doc.add_paragraph('')
    theme_r.add_run(themas[theme]).italic = True
    theme_r.alignment = 1
    question_print = doc.add_paragraph()
    if per > len(q)-1:
        per = 0

    question_print.add_run(q[per])
    


    doc.add_page_break()
    per = per + 1 






# add a run i.e, style like 
# bold, italic, underline, etc.
# doc_para.add_run('hey there, bold here').bold = True
# doc_para.add_run(', and ')
# doc_para.add_run('these words are italic').italic = True
  
# add a page break to start a new page
# doc.add_page_break()
  
# add a heading of level 2
# doc.add_heading('Heading level 2', 2)


mydir = '.'
myfile = f'questions{theme}.docx'
# now save the document to a location
doc.save(os.path.join(mydir, myfile))